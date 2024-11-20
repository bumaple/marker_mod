import os
import io
import argparse
import re
import time
import json

from datetime import datetime
from typing import Dict, Tuple

from loguru import logger

from marker.database.pdf_data_operator import PDFDataOperator
from marker.http.http_utils import HttpUtils
from marker.logger import set_logru
from marker.config_read import Config
from marker.output import save_html_markdown_json, get_subfolder_path
from marker.docx.data.docx_content_data import DocxData, DocxChapterData, DocxParagraphData

from docx import Document
from PIL import Image
import hashlib
import shutil


def get_image_dimensions(image_data):
    """获取图片尺寸"""
    try:
        with Image.open(io.BytesIO(image_data)) as img:
            return img.size
    except Exception:
        return None, None


def get_safe_filename(original_name, image_data):
    """生成安全的文件名，基于图片内容的哈希值"""
    hash_obj = hashlib.md5(image_data)
    hash_value = hash_obj.hexdigest()[:8]
    extension = os.path.splitext(original_name)[1] if original_name else '.png'
    return f"img_{hash_value}{extension}"


def find_image_rids_in_paragraph(paragraph):
    """在段落中查找图片的 rId"""
    rids = []
    for run in paragraph.runs:
        if not hasattr(run, '_element'):
            continue
        drawings = run.element.findall('.//wp:inline//a:blip',
                                       {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if drawings is None:
            continue
        for drawing in drawings:
            rid = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rid:
                rids.append(rid)
    return rids


def extract_images(doc, output_folder):
    """从Word文档中提取图片并保存到指定文件夹"""
    image_map = {}
    # 处理文档级别的关系
    rels = doc.part.rels
    for rel in rels.values():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                original_name = rel.target_part._path if hasattr(rel.target_part, '_path') else ''
                # 生成安全的文件名
                image_filename = get_safe_filename(original_name, image_data)
                image_path = os.path.join(output_folder, image_filename)
                # 获取图片尺寸
                width, height = get_image_dimensions(image_data)
                # 保存图片
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                # 存储图片信息
                image_map[rel.rId] = {
                    'filename': image_filename,
                    'width': width,
                    'height': height,
                    'original_name': original_name
                }
            except Exception as e:
                print(f"Warning: Failed to process image {rel.rId}: {str(e)}")
                continue
    return image_map


STANDALONE_NUMBER_PATTERNS = [
    # r'^[A-Za-z](\.\d+)*\.$',  # 字母数字组合带点: A.1., A.1.1.
    (r'^[A-Za-z](\.\d+)+$', 0),  # 字母数字组合: A.1, A.1.1
    # r'^[A-Za-z]\.$',  # 带点的字母: A., B., C.
    (r'^[A-Za-z]$', 0),  # 单个字母: A, B, C
    # r'^\d+(\.\d+)*\.$',  # 多级数字带点: 1.1., 1.1.1.
    (r'^\d+(\.\d+)+$', 0),  # 多级数字: 1.1, 1.1.1
    # r'^\d+\.$',  # 带点的数字: 1., 2., 3.
    (r'^\d+$', 0),  # 纯数字: 1, 2, 3
]

NUMBERING_PATTERNS = [
    # r'(?<!\S)[A-Z](?:\.\d+)+\.\s',  # 1. 单个字母数字组合带点: A.1., A.1.1.
    (r'(?<!\S)[A-Z](?:\.\d+)+\s', 0),  # 2. 单个字母数字组合: A.1, A.1.1
    # r'(?<!\S)[A-Z]\.\s',  # 3. 单个字母带点: A., B., C.
    (r'(?<!\S)附\s?录\s?[A-Z](?!\S)', 2),  # 4. 附录带单个字母组合: 附录A, 附录B, 附录C
    # r'(?<!\S)\d+(?:\.\d+)+\.\s',  # 5. 多级数字带点: 1.1., 1.1.1.
    (r'(?<!\S)\d+(?:\.\d+)+\s', 0),  # 6. 多级数字: 1.1, 1.1.1
    # r'(?<!\S)\d+\.\s',  # 7. 带点的数字: 1., 2., 3.
    (r'(?<!\S)\d+(?!\S)', 0),  # 8. 纯数字: 1, 2, 3
    (r'(?<!\S)前\s?言(?!\S)', 1),  # 9. 特殊章节: 前言, 前 言
]

REPLACE_PATTERNS = [
    # 单个字母数字组合带点: A. 1., A.1. 1.
    (r'([A-Z])\s*\.\s*(\d+(?:\s*\.\s*\d+)*)\s*\.', r'\1.\2.'),
    # 单个字母数字组合: A .1, A.1 . 1
    (r'([A-Z])\s*(\d+(?:\s*\.\s*\d+)*)', r'\1\2'),
    # 单个字母带点: A ., B., C.
    (r'([A-Z])\s*\.', r'\1.'),
    # 附录带单个字母组合: 附 录A, 附录 B, 附 录 C
    (r'附\s*录\s*([A-Z])', r'附录\1'),
    # 多级数字: 1 .1, 1. 1. 1
    (r'(\d+)(?:\s*\.\s*(\d+))+', lambda m: '.'.join(m.group().split())),
    # 多级数字带点: 1 . 1., 1. 1 .1 .
    (r'(\d+(?:\s*\.\s*\d+)*)\s*\.', r'\1.'),
    # 带点的数字: 1 ., 2 ., 3.
    (r'(\d+)\s*\.', r'\1.'),
    # 特殊章节: 前 言
    (r'前\s*言', '前言'),
    # 处理汉字短语中的空格: 中 华 人民 共 和国
    (r'\s*'.join(['中', '华', '人', '民', '共', '和', '国']), '中华人民共和国'),
    # 处理汉字短语中的空格:
    (r'\s*'.join(['国', '家', '标', '准']), '国家标准'),
    # 处理汉字短语中的空格:
    (r'\s*'.join(['地', '方', '标', '准']), '地方标准'),
    # 处理汉字短语中的空格:
    (r'\s*'.join(['行', '业', '标', '准']), '行业标准'),
    # 替换破折号为连字符
    (r'—', '-'),
    # 替换连续多个空格为单个空格
    (r'\s{2,}', ' '),
    # 特殊文本: 表 1
    (r'表\s*(\d+)', r'表\1.'),
    # 特殊文本: 图 1
    (r'图\s*(\d+)', r'图\1.'),
]


def clean_text(text):
    text = text.strip()
    if len(text) == 0:
        return text
    for pattern, repl in REPLACE_PATTERNS:
        text = re.sub(pattern, repl, text)
    return text


def remove_chinese(text):
    pattern = re.compile(r'[\u4e00-\u9fa5]+')
    return re.sub(pattern, '', text)


def is_chemical_formula(text):
    """
    识别文本是否为化学式
    匹配模式：字母+数字的组合，可能包含括号和特殊字符
    """
    # 更严格的化学式模式：
    # 1. 必须包含至少一个大写字母和数字的组合
    # 2. 可能包含括号和其他化学式常见字符
    if not text:
        return False

    # 检查是否符合化学式的基本特征
    has_caps_number = bool(re.search(r'[A-Z][a-z]?\d', text))
    potential_formula = bool(re.match(r'^[A-Z][a-z0-9().\-+=]*$', text))

    # 常见的化学元素列表（部分示例）
    common_elements = {'H', 'C', 'N', 'O', 'F', 'Cl', 'Br', 'I', 'Na', 'Fe', 'Cu', 'Au', 'Ag'}

    # 检查第一个元素是否是常见化学元素
    first_element = re.match(r'([A-Z][a-z]?)', text)
    has_common_element = first_element and first_element.group(1) in common_elements

    return has_caps_number and potential_formula and has_common_element


def split_chemical_formula(formula):
    """
    将化学式拆分为基本元素和数字
    返回一个列表，每个元素是(元素符号, 数字)的元组
    """
    # 使用正则表达式匹配元素符号和数字
    pattern = r'([A-Z][a-z]?)(\d*)'
    matches = re.findall(pattern, formula)
    return [(symbol, num if num else '1') for symbol, num in matches]


HTML_FONT_SIZE_DEFAULT = 12
HTML_FONT_SIZE_MIN = 10.5


def get_font_size(run):
    """
    获取运行文本的字体大小
    如果没有明确设置，返回默认值
    """
    try:
        if run.font.size is not None:
            # 字体大小通常以twip为单位（1/20点）
            if run.font.size.pt >= HTML_FONT_SIZE_MIN:
                return HTML_FONT_SIZE_DEFAULT
            return run.font.size.pt
        return HTML_FONT_SIZE_DEFAULT  # 默认字体大小
    except AttributeError:
        return HTML_FONT_SIZE_DEFAULT


def is_mergeable_separator(item):
    """检查元素是否为可以合并表格的分隔符（分页符、分节符或回车换行）"""
    if item.tag.endswith('sectPr') or item.tag.endswith('br') or item.tag.endswith('p'):
        # 检查段落是否为空白段落（如回车换行）以决定是否可以合并
        if item.tag.endswith('p'):
            if len(item.text.strip()) == 0:  # 检查段落是否没有文本
                return True
        else:
            return True
    return False


def should_merge_tables(doc_body, current_index):
    """
    检查当前表格是否应该与下一个表格合并
    返回下一个表格的索引，如果不应该合并则返回None
    """
    if current_index >= len(doc_body) - 1:  # 如果是最后一个元素，不需要合并
        return None

    # 从当前表格后面一个元素开始查找
    index = current_index + 1
    while index < len(doc_body):
        item = doc_body[index]
        # 如果找到另一个表格
        if item.tag.endswith('tbl'):
            # 检查从当前表格到这个表格之间的所有元素是否都是可合并的分隔符
            can_merge = all(
                is_mergeable_separator(doc_body[i])
                for i in range(current_index + 1, index)
            )
            return index if can_merge else None
        # 如果遇到非可合并的元素，停止查找
        elif not is_mergeable_separator(item):
            return None
        index += 1
    return None


def get_merged_cell_info(table) -> dict:
    """获取表格中所有单元格的合并信息"""
    merged_cells = {}
    row_cells, col_cells = [], []

    # 获取行列单元格信息
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(cell)
        row_cells.append(cells)

    for col in table.columns:
        cells = []
        for cell in col.cells:
            cells.append(cell)
        col_cells.append(cells)

    # 计算每个单元格的合并信息
    for i in range(len(table.rows)):
        for j in range(len(table.columns)):
            cell = table.rows[i].cells[j]
            cell_key = (i, j)

            # 检查是否已经处理过这个单元格
            if cell_key not in merged_cells:
                col_counts = row_cells[i].count(row_cells[i][j])  # 水平合并数
                row_counts = col_cells[j].count(col_cells[j][i])  # 垂直合并数

                # 存储合并信息
                merged_cells[cell_key] = {
                    'cell': cell,
                    'colspan': col_counts,
                    'rowspan': row_counts,
                    'content': get_cell_content(cell)
                }

                # 标记被合并的单元格
                if col_counts > 1 or row_counts > 1:
                    for ri in range(i, i + row_counts):
                        for ci in range(j, j + col_counts):
                            if (ri, ci) != cell_key:
                                merged_cells[(ri, ci)] = None

    return merged_cells


def get_cell_content(cell):
    """获取单元格的完整内容，包括段落和格式"""
    content = []
    for paragraph in cell.paragraphs:
        para_content = []
        for run in paragraph.runs:
            para_content.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline
            })
        content.append(para_content)
    return content


def apply_cell_content(cell, content):
    """将保存的内容应用到单元格"""
    # 清除现有内容
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # 添加新内容
    for paragraph_content in content:
        paragraph = cell.add_paragraph()
        for run_info in paragraph_content:
            run = paragraph.add_run(run_info['text'])
            run.bold = run_info['bold']
            run.italic = run_info['italic']
            run.underline = run_info['underline']


def merge_tables(table1, table2):
    """
    合并两个表格，保持单元格合并状态
    """
    # 获取两个表格的列数
    num_cols1 = len(table1.columns)
    num_cols2 = len(table2.columns)

    # 确定使用哪个表格作为基础
    if num_cols1 >= num_cols2:
        base_table = table1
        added_table = table2
        max_cols = num_cols1
    else:
        base_table = table2
        added_table = table1
        max_cols = num_cols2

    # 获取要添加的表格的合并信息
    added_merged_cells = get_merged_cell_info(added_table)

    # 获取基础表格的当前行数
    base_rows = len(base_table.rows)

    # 创建足够的行来容纳新内容
    rows_needed = len(added_table.rows)
    for _ in range(rows_needed):
        new_row = base_table.add_row()
        # 确保新行有正确数量的单元格
        while len(new_row.cells) < max_cols:
            new_row._tr.append_cell()

    # 处理单元格合并和内容
    for i in range(len(added_table.rows)):
        row_cells_processed = set()

        for j in range(min(len(added_table.columns), max_cols)):
            cell_key = (i, j)

            # 跳过已经处理过的单元格（被合并的单元格）
            if cell_key in row_cells_processed:
                continue

            merge_info = added_merged_cells.get(cell_key)
            if merge_info is None:
                continue

            # 获取目标单元格
            current_row_idx = base_rows + i
            try:
                target_cell = base_table.rows[current_row_idx].cells[j]
            except IndexError:
                continue  # 跳过超出范围的单元格

            # 应用内容
            apply_cell_content(target_cell, merge_info['content'])

            # 处理单元格合并
            if merge_info['colspan'] > 1 or merge_info['rowspan'] > 1:
                # 计算合并范围（确保不超出表格边界）
                start_row = current_row_idx
                end_row = min(start_row + merge_info['rowspan'] - 1, len(base_table.rows) - 1)
                start_col = j
                end_col = min(j + merge_info['colspan'] - 1, max_cols - 1)

                # 标记被合并的单元格
                for ri in range(i, i + merge_info['rowspan']):
                    for ci in range(j, j + merge_info['colspan']):
                        if ci < max_cols:  # 确保不超出列范围
                            row_cells_processed.add((ri, ci))

                # 只有当合并范围有效时才执行合并
                if start_row < end_row or start_col < end_col:
                    try:
                        target_cell = base_table.cell(start_row, start_col)
                        merged_cell = base_table.cell(end_row, end_col)
                        target_cell.merge(merged_cell)
                    except IndexError:
                        continue  # 如果合并失败，继续处理下一个单元格

    return base_table


def get_heading_level(paragraph) -> Tuple[int, int]:
    """
    检查段落是否是标题样式，如果是则返回标题级别（1-6），否则返回None
    """
    if not hasattr(paragraph, 'style') or not paragraph.style:
        return 0, 0
    style_name = paragraph.style.name.lower() if paragraph.style.name else ''
    if 'heading' in style_name or 'title' in style_name:
        match = re.search(r'\d+', style_name)
        if match:
            level = int(match.group())
            return min(max(level, 1), 6), 0
        elif 'title' in style_name:
            return 1, 0
    return 0, 0


def get_number_heading_level(text: str) -> Tuple[int, int]:
    """
    根据编号格式判断标题级别
    返回对应的标题级别（1-6），如果不是有效的编号格式则返回None
    """
    text = text.strip()
    for pattern, level_flag in NUMBERING_PATTERNS:
        match = re.match(pattern, text)
        if match:
            if level_flag != 0:
                return 1, level_flag
            # 计算级别
            level = len(re.findall(r'\d+|\w', match.group()))
            return min(level, 6), level_flag
    for pattern, level_flag in STANDALONE_NUMBER_PATTERNS:
        match = re.match(pattern, text)
        if match:
            if level_flag != 0:
                return 1, level_flag
            # 计算级别
            level = len(re.findall(r'\d+|\w', match.group()))
            return min(level, 6), level_flag
    return 0, 0


def get_combined_heading_level(paragraph, text) -> Tuple[int, int]:
    """
    综合考虑段落样式和编号格式来确定标题级别
    返回最终的标题级别（1-6）或 0
    """
    style_level, style_level_flag = get_heading_level(paragraph)
    number_level, number_level_flag = get_number_heading_level(text)

    if style_level == 0 and number_level == 0:
        return 0, 0
    if style_level == 0:
        return number_level, number_level_flag
    if number_level == 0:
        return style_level, style_level_flag
    return min(style_level, number_level), min(style_level_flag, number_level_flag)


def convert_table_to_html(table, image_map, is_styled=True) -> Tuple[str, dict]:
    """转换表格为HTML，处理单元格合并"""
    # 获取表格的合并信息
    merged_cells = get_merged_cell_info(table)

    table_text = '<table>'

    for i in range(len(table.rows)):
        table_text = table_text + '<tr>\n'
        cells_processed = set()

        for j in range(len(table.columns)):
            cell_key = (i, j)

            # 跳过已经处理过的单元格（被合并的单元格）
            if cell_key in cells_processed:
                continue

            merge_info = merged_cells.get(cell_key)
            if merge_info is None:
                continue

            cell = merge_info['cell']
            colspan = merge_info['colspan']
            rowspan = merge_info['rowspan']

            # 处理单元格内容
            cell_html = []
            for paragraph in cell.paragraphs:
                para_info = convert_paragraph_to_html(paragraph, image_map)
                if para_info['text']:
                    cell_html.append(para_info['text'])
                if para_info['has_images']:
                    cell_html.extend(para_info['html'])

            # 合并所有段落内容
            cell_content = '<br>'.join(cell_html) if cell_html else '&nbsp;'

            # 添加合并属性
            merge_attrs = []
            if colspan > 1:
                merge_attrs.append(f'colspan="{colspan}"')
            if rowspan > 1:
                merge_attrs.append(f'rowspan="{rowspan}"')

            # 生成单元格HTML
            attrs = ' '.join(merge_attrs)
            table_text = table_text + f'<td {attrs}>{cell_content}</td>'

            # 标记被合并的单元格
            for ri in range(i, i + rowspan):
                for ci in range(j, j + colspan):
                    cells_processed.add((ri, ci))

        table_text = table_text + '\n</tr>\n'

    table_text = table_text + '</table>\n'
    return table_text, merged_cells


def convert_paragraph_to_html(paragraph, image_map):
    """将段落转换为HTML格式，并处理独立编号合并到下一行的逻辑"""
    html = []
    text_content = ''
    has_images = False
    # number_text = None

    # 存储run信息，包括文本和字体大小
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'size': get_font_size(run),
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'sup': run.font.superscript,
            'sub': run.font.subscript
        })

    # 处理文本内容和样式
    i = 0
    while i < len(runs_info):
        run_info = runs_info[i]
        current_text = run_info['text']
        sup = run_info['sup']
        sub = run_info['sub']

        if not current_text:
            i += 1
            continue

        if sup:
            current_text = f'<sup>{current_text}</sup>'
        elif sub:
            current_text = f'<sub>{current_text}</sub>'
        else:
            # 检查字体大小差异
            if i < len(runs_info) - 1:
                next_run = runs_info[i + 1]
                size_diff = abs(run_info['size'] - next_run['size'])
                if size_diff > 4:
                    if run_info['size'] < next_run['size']:
                        current_text = f'<sub>{current_text}</sub>'

        # 处理化学式
        chemical_pattern = r'(?:^|\s|[(（])([A-Z][a-z]?\d*(?:[A-Z][a-z]?\d*)*(?:-\d*)?|C\d+)(?=\s|$|[,，.。)）])'
        parts = []
        last_end = 0

        for match in re.finditer(chemical_pattern, current_text):
            # 添加化学式前的文本
            parts.append(current_text[last_end:match.start(1)])
            # 处理化学式
            chemical = match.group(1)
            if is_chemical_formula(chemical):
                formula_parts = split_chemical_formula(chemical)
                formula_html = []
                for element, number in formula_parts:
                    if number != '1':
                        formula_html.append(f'{element}<sub>{number}</sub>')
                    else:
                        formula_html.append(element)
                parts.append(''.join(formula_html))
            else:
                parts.append(chemical)
            last_end = match.end(1)

        # 添加剩余文本
        parts.append(current_text[last_end:])
        # 合并处理后的文本
        processed_text = ''.join(parts)

        text_content += processed_text
        i += 1

    # 处理图片
    rids = find_image_rids_in_paragraph(paragraph)
    for rid in rids:
        if rid in image_map:
            has_images = True
            img_info = image_map[rid]
            img_attrs = []
            img_path = f"images/{img_info['filename']}"
            img_attrs.append(f'src="{img_path}"')
            alt_text = os.path.splitext(img_info['original_name'])[0] or 'image'
            img_attrs.append(f'alt="{alt_text}"')
            if img_info['width'] and img_info['height']:
                img_attrs.append(f'width="{img_info["width"]}"')
                img_attrs.append(f'height="{img_info["height"]}"')
            img_attrs.append('style="max-width: 100%; height: auto;"')
            html.append(f'<img {" ".join(img_attrs)}>')

    header_level = 0
    header_level_flag = 0
    text_content = clean_text(text_content)
    if len(text_content) > 0:
        header_level, header_level_flag = get_combined_heading_level(paragraph, text_content)

    # 返回包含段落信息的字典
    return {
        'text': text_content,  # text_content 已经包含规范化的空格
        'has_images': has_images,
        'html': html,
        'header_level': header_level,
        'header_level_flag': header_level_flag,
    }


def process_doc_body(doc):
    """处理文档主体，合并相邻的可合并表格"""
    processed_items = []
    i = 0
    while i < len(doc.element.body):
        item = doc.element.body[i]

        if item.tag.endswith('tbl'):
            current_table = [t for t in doc.tables if t._element == item][0]
            next_table_index = should_merge_tables(doc.element.body, i)

            # 如果找到可以合并的下一个表格
            while next_table_index is not None:
                next_table = [t for t in doc.tables if t._element == doc.element.body[next_table_index]][0]
                current_table = merge_tables(current_table, next_table)
                i = next_table_index  # 更新索引到最后合并的表格位置
                # 继续查找下一个可合并的表格
                next_table_index = should_merge_tables(doc.element.body, i)

            processed_items.append(('table', current_table))
        elif not (item.tag.endswith('sectPr') or item.tag.endswith('br')):
            # 保存非分隔符的其他内容
            processed_items.append(('other', item))
        i += 1

    return processed_items


def convert_word_markdown_json(config_data, file_path, file_name, file_to_convert, metadata) -> Tuple[str, str, DocxData, int, Dict]:
    """将Word文档转换为HTML，处理编号合并逻辑和表格合并"""
    # 创建images文件夹
    images_dir = os.path.join(file_path, 'images')
    if os.path.exists(images_dir):
        shutil.rmtree(images_dir)  # 清理旧的图片目录
    os.makedirs(images_dir)

    # 加载Word文档
    doc = Document(file_to_convert)
    # 提取图片
    image_map = extract_images(doc, images_dir)

    # 生成HTML内容
    html_content = ['''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
table { margin-bottom: 15px; border-collapse: collapse; width: 100%; }
td { border: 0.5pt solid black; padding: 5px; vertical-align: top; }''',
                    'p { margin: 10px 0; font-size: ' + str(HTML_FONT_SIZE_DEFAULT) + 'pt; }',
                    '''img { max-width: 100%; height: auto; display: block; margin: 10px 0; }
                    h1, h2, h3, h4, h5, h6 { margin: 20px 0 10px 0; font-weight: bold; }
                    h1 { font-size: 14pt; }
                    h2 { font-size: 12pt; }
                    h3 { font-size: 12pt; }
                    h4 { font-size: 12pt; }
                    h5 { font-size: 12pt; }
                    h6 { font-size: 12pt; }
                    h1:not(:empty), h2:not(:empty), h3:not(:empty), h4:not(:empty), h5:not(:empty), h6:not(:empty) { padding-left: 10px; border-left: 4px solid #333; line-height: 1.4; }
                    </style>
                    </head>
                    <body>''']

    markdown_content = []

    docx_data = DocxData(fileName=file_name)

    # 处理文档主体，合并表格
    processed_items = process_doc_body(doc)

    chapter_seq = 1
    docx_chapter_data = DocxChapterData(sn='封面', title='', level=0, seq=chapter_seq)
    docx_data.add_chapter(docx_chapter_data)

    for item_type, item in processed_items:
        if item_type == 'table':
            try:
                convert_html, merged_cells_dict = convert_table_to_html(item, image_map)
                html_content.append(convert_html)

                ayalyzing_content = analyzing_docx_table(config_data, convert_html)
                markdown_content.append(ayalyzing_content)

                if len(ayalyzing_content) > 0:
                    docx_paragraph_data = DocxParagraphData(content=convert_html, ayalyzing_content=ayalyzing_content, type=DocxParagraphData.TYPE_TABLE)
                    docx_chapter_data.add_paragraph(docx_paragraph_data)
            except Exception as e:
                logger.error(f"表格处理错误: {e}")
                return '', '', docx_data, 9, metadata
        else:
            if item.tag.endswith('p'):
                paragraph = [p for p in doc.paragraphs if p._element == item][0]
                para_info = convert_paragraph_to_html(paragraph, image_map)
                # 标题级别
                header_level = para_info['header_level']
                header_level_flag = para_info['header_level_flag']
                if para_info['text']:
                    if header_level > 0:
                        # 获取章节编号
                        title_content = ''.join(para_info['text'])
                        title_list = title_content.split(" ")
                        if len(title_list) > 0:
                            sn_text = title_list[0].replace(' ', '')
                        else:
                            sn_text = title_content.replace(' ', '')
                        title_full_content = sn_text
                        title_text = ''
                        if len(title_list) > 1:
                            title_text = ' '.join(title_list[1:])
                            title_full_content += f" {title_text}"

                        html_content.append(f"<h{header_level}>{title_full_content}</h{header_level}>")
                        markdown_content.append(f"{header_level * '#'} {title_full_content}")

                        if header_level_flag == 2:
                            title_key = remove_chinese(sn_text)
                        else:
                            title_key = sn_text

                        docx_chapter_data, chapter_seq = docx_data.find_chapter_by_key(key=title_key, sn=sn_text, title=title_text, level=header_level, seq=chapter_seq)
                    else:
                        html_content.append(f'<p>{para_info["text"]}</p>')
                        markdown_content.append(f'{para_info["text"]}')

                        json_text = ''.join(para_info["text"])
                        if len(json_text) > 0:
                            docx_paragraph_data = DocxParagraphData(content=json_text, type=DocxParagraphData.TYPE_TEXT)
                            docx_chapter_data.add_paragraph(docx_paragraph_data)
                if para_info['has_images']:
                    html_content.extend(para_info['html'])
                    markdown_content.extend(para_info['html'])

                    json_text = ''.join(para_info["html"])
                    if len(json_text) > 0:
                        docx_paragraph_data = DocxParagraphData(content=json_text, type=DocxParagraphData.TYPE_IMAGE)
                        docx_chapter_data.add_paragraph(docx_paragraph_data)

    html_content.append('</body>\n</html>')

    return '\n'.join(html_content), '\n'.join(markdown_content), docx_data.to_json(), 1, metadata


def convert_handler(config_data, pdf_data_opt, data_source, max_files, metadata_list, files) -> Tuple[int, str, list]:
    if len(files) == 0:
        return 0, '待处理文件为空', []

    start_time = datetime.now()

    # 处理最大文件数
    if data_source == 'path' and max_files:
        files_to_convert = files[:max_files]
    else:
        files_to_convert = files

    files_number = len(files_to_convert)
    success_number = 0

    # 执行过程开始
    return_html_files = []
    for idx, file_to_convert in enumerate(files_to_convert):
        file_name = os.path.basename(file_to_convert)
        metadata = metadata_list.get(file_name)

        if data_source == 'db':
            if 'record_id' not in metadata:
                log_info = f"MetaData 中不存在 id！ {metadata}"
                logger.error(log_info)
                return 0, f"MetaData 中不存在 id！ {metadata}", []

            if 'ocr_types' not in metadata:
                log_info = f"MetaData 中不存在 ocr_types！ {metadata}"
                logger.error(log_info)
                return 0, f"MetaData 中不存在 ocr_types！ {metadata}", []

            record_id = metadata['record_id']
            ocr_type_strs = metadata['ocr_types']
        else:
            record_id = -1
            ocr_type_strs = '00'

        if 'out_folder' not in metadata:
            out_folder = os.path.dirname(file_to_convert)
        else:
            out_folder = metadata['out_path']

        subfolder_path = get_subfolder_path(out_folder, file_name, ocr_type_strs)

        try:
            html_content, markdown_content, json_content, result_code, out_metadata = convert_word_markdown_json(config_data,
                subfolder_path, file_name, file_to_convert, metadata)
            if len(html_content) > 0 or len(markdown_content) > 0 or json_content is not None:
                html_path, html_file, markdown_file, json_file, metadata_file = save_html_markdown_json(out_folder,
                                                                                                        file_name,
                                                                                                        html_content,
                                                                                                        markdown_content,
                                                                                                        json_content,
                                                                                                        out_metadata,
                                                                                                        ocr_type_strs,
                                                                                                        subfolder_path)
                return_html_files.append(
                    {'path': html_path, 'html_file': html_file, 'markdown_file': markdown_file, 'json_file': json_file, 'metadata_file': metadata_file})
                logger.info(f"写入文件成功! 第{idx + 1}个 文件名：{file_name} 保存路径：{html_path}")

                success_number += 1

                if data_source == 'db' and record_id != '':
                    json_file_name = os.path.basename(json_file)
                    result_num = pdf_data_opt.update_pri_fix_file(record_id, json_file, json_file_name)
                    if result_num > 0:
                        pdf_data_opt.insert_pri_docx_handler(record_id)
            else:
                if result_code == 9:
                    return 0, '表格格式错误', []
        except Exception as e:
            logger.error(f"处理文件失败！第{idx + 1}个 文件名：{file_name} 路径：{out_folder} 错误：{e}")

    # 执行过程结束

    end_time = datetime.now()
    # 计算实际执行的时间
    execution_time = end_time - start_time
    execution_seconds = execution_time.total_seconds()

    # 将执行时间转换为时分秒格式
    hours, remainder = divmod(execution_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)

    average_time = round(execution_seconds / len(files_to_convert))

    log_info = f" * * * * * 处理完成！文件数：{files_number} [成功 {success_number}，失败 {files_number - success_number}]。完成时间：{end_time.strftime('%Y-%m-%d %H:%M:%S')}。总处理时间：{int(hours)} 小时 {int(minutes)} 分 {int(seconds)} 秒，处理速度：{average_time} 秒/个"
    logger.info(log_info)
    return len(files_to_convert), '处理完成', return_html_files


def get_data_from_db(pdf_data_opt, batch_number) -> Tuple[int, list, dict]:
    metadata_list = {}
    files = []

    records = pdf_data_opt.query_need_docx(batch_number)
    if len(records) <= 0:
        log_info = f"没有需要处理的数据！"
        logger.info(log_info)

    # 循环输出查询结果
    for row in records:
        record_id = row['id']
        word_file = row['word_file']
        word_title = row['title']
        # json_file = row['json_file']
        # json_file_name = row['json_file_name']
        file_name = os.path.basename(word_file)
        out_folder = os.path.dirname(word_file)
        if file_name.endswith('.doc') or file_name.endswith('.docx'):
            metadata_list[file_name] = {"out_path": out_folder,
                                       "record_id": record_id, "title": word_title, "ocr_types": '00'}
            if os.path.isfile(word_file):
                files.append(word_file)
            else:
                logger.warning(f"文件不存在：{word_file}")
        else:
            logger.warning(f"文件不是doc或docx格式，跳过：{word_file}")
    return 1, files, metadata_list


def get_data_from_path(metadata_file_arg, in_folder_arg, out_folder_arg, ocr_types) -> Tuple[int, list, dict]:
    if metadata_file_arg:
        metadata_file = os.path.abspath(metadata_file_arg)
        with open(metadata_file, "r") as f:
            metadata_list = json.load(f)
    else:
        metadata_list = {}

    if in_folder_arg is None:
        logger.error(f"缺少 --in_folder 参数！")
        return 0, [], {}

    if out_folder_arg is None:
        if os.path.isfile(in_folder_arg):
            out_folder_arg = os.path.dirname(in_folder_arg)
        else:
            out_folder_arg = in_folder_arg

    in_folder = os.path.abspath(in_folder_arg)
    out_folder = os.path.abspath(out_folder_arg)

    if os.path.isfile(in_folder):
        files = [in_folder]
    else:
        os.makedirs(out_folder, exist_ok=True)
        files = [os.path.join(in_folder, f) for f in os.listdir(in_folder)]
        files = [f for f in files if os.path.isfile(f) and f.endswith('.doc') or f.endswith('.docx')]
    for file in files:
        file_name = os.path.basename(file)
        metadata_list[file_name] = {"out_path": out_folder, "record_id": '',
                                    "ocr_types": ocr_types}
    return 1, files, metadata_list


def check_files_status(pdf_data_opt, max_files_arg) -> Tuple[int, list]:
    records = pdf_data_opt.query_sub_finish_fix(max_files_arg)

    error_files = []
    # 循环输出查询结果
    for row in records:
        # record_id = row['ID']
        md_file_path = row['MD_FILE_DIR']
        md_file_name = row['MD_FILE_NAME']
        md_file = os.path.join(md_file_path, md_file_name)
        if md_file.endswith('.md'):
            if not os.path.isfile(md_file):
                error_files.append(md_file)
        else:
            error_files.append(md_file)
    return len(records), error_files


def analyzing_docx_table(config_data: Config, table_html: str) -> str:
    prompt_head = "你是处理HTML格式的专家，任务是将HTML表格处理为符合要求的文本并按照要求输出。请遵循以下要求：\n" + \
                  "1.不添加原文中不存在的任何新信息；\n" + \
                  "2.理解HTML表格，整合为一段文字内容；\n" + \
                  "3.只回复符合格式要求的文本，不添加任何引言、解释或元数据。\n"

    http_utils = HttpUtils(config_data)
    resp_content, resp_result = http_utils.request_openai_api(prompt_head=prompt_head, text_content=table_html)
    return resp_content


def main():
    parser = argparse.ArgumentParser(description="转化docx文件为html及json.")
    parser.add_argument("--in_folder", help="Input folder with files.")
    parser.add_argument("--out_folder", help="Output folder with files.")
    parser.add_argument("--max", type=int, default=0, help="Maximum number of files to fix")
    parser.add_argument("--metadata_file", type=str, default=None, help="Metadata json file to use for filtering")

    parser.add_argument("--config_file", default='config.ini', help="config file.")
    # 增加操作类型，convert：识别转化PDF check：检查转化效果
    parser.add_argument("--run_type", default='convert', help="run type type (convert or check)")
    parser.add_argument("--ocr_types", type=str, default='00', help="OCR type (40:word文件)")

    args = parser.parse_args()

    start_time = datetime.now()

    # 从配置文件中读取数据库配置
    in_folder_arg = args.in_folder
    out_folder_arg = args.out_folder
    max_files_arg = args.max
    metadata_file_arg = args.metadata_file
    config_file_arg = args.config_file
    run_type_arg = args.run_type.lower()
    ocr_types_arg = args.ocr_types

    config = Config(config_file_arg)

    data_source = config.get_sys_param('data_source')
    if data_source is None:
        data_source = ''
    data_source = data_source.lower()

    batch_number = config.get_sys_param('batch_number', int)
    if batch_number == 0:
        batch_number = 100

    sleep_minute = config.get_sys_param('sleep_minute', int)
    if sleep_minute == 0:
        sleep_minute = 10

    log_level = config.get_sys_param('log_level')
    if log_level is not None:
        set_logru(log_level=log_level)
    else:
        set_logru()

    if run_type_arg == 'convert':
        if data_source == 'db':
            pdf_data_opt = PDFDataOperator(config_file_arg)
            while True:
                result_code, files, metadata_list = get_data_from_db(pdf_data_opt, batch_number)
                if result_code == 0:
                    return

                result_code, result_msg, out_file = convert_handler(config, pdf_data_opt, data_source, max_files_arg,
                                                                    metadata_list, files)
                if result_code == 0:
                    return
                elif result_code > 0:
                    time.sleep(60)
                else:
                    time.sleep(sleep_minute * 60)
        elif data_source == 'path':
            result_code, files, metadata_list = get_data_from_path(metadata_file_arg, in_folder_arg, out_folder_arg,
                                                                   ocr_types_arg)
            if result_code == 0:
                return

            result_code, result_msg, out_file = convert_handler(config,None, data_source, max_files_arg, metadata_list,
                                                                files)
            if result_code == 0:
                return
        else:
            log_info = f"不支持的data_source参数！{data_source}"
            logger.error(log_info)
            return
    else:
        pdf_data_opt = PDFDataOperator(config_file_arg)
        record_cnt, error_files = check_files_status(pdf_data_opt, max_files_arg)
        log_info = f" * * * * * {run_type_arg.capitalize()}ed 文件数：{record_cnt}。{len(error_files)} 个文件不存在！"
        logger.info(log_info)
        if len(error_files) > 0:
            for error_file in error_files:
                logger.info(error_file)


if __name__ == "__main__":
    main()
